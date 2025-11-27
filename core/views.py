from django.shortcuts import render, redirect, get_object_or_404
import requests
from .models import *
from django.http import JsonResponse, HttpResponseBadRequest, HttpResponse
from .forms import *
from rest_framework import viewsets
from .serializers import *
from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.serializers import serialize
from decimal import Decimal
from django.views.decorators.csrf import csrf_exempt
import docx
from io import BytesIO
from functools import wraps
from django.http import HttpResponseForbidden
from urllib.parse import quote_plus
from django.shortcuts import get_object_or_404
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from django.db.models import Avg, Exists, OuterRef
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill
from openpyxl.utils import get_column_letter
from rest_framework.response import Response
from rest_framework import status
from rest_framework.decorators import api_view
from django.contrib.auth.hashers import check_password
from django.contrib.auth import authenticate, login as auth_login
from django.contrib.auth.views import LogoutView 


def role_required(roles):
    def decorator(view_func):
        @wraps(view_func)
        def _wrapped_view(request, *args, **kwargs):
            if request.user.is_authenticated and request.user.idRol and request.user.idRol.nombreRol in roles:
                return view_func(request, *args, **kwargs)
            return HttpResponseForbidden("No tienes permiso para acceder a esta página")
        return _wrapped_view
    return decorator

# ViewSets para modelos relacionados al usuario
class RolUsuarioViewSet(viewsets.ModelViewSet):
    queryset = rolUsuario.objects.all()
    serializer_class = RolUsuarioSerializer

class RegionViewSet(viewsets.ModelViewSet):
    queryset = Region.objects.all()
    serializer_class = RegionSerializer

class ComunaViewSet(viewsets.ModelViewSet):
    queryset = Comuna.objects.all()
    serializer_class = ComunaSerializer

class UsuarioCustomViewSet(viewsets.ModelViewSet):
    queryset = UsuarioCustom.objects.all()
    serializer_class = UsuarioCustomSerializer

class PreferenciasUsuarioViewSet(viewsets.ModelViewSet):
    queryset = PreferenciasUsuario.objects.all()
    serializer_class = PreferenciasUsuarioSerializer

# ViewSets para modelos relacionados al taller
class TallerViewSet(viewsets.ModelViewSet):
    queryset = Taller.objects.all()
    serializer_class = TallerSerializer

class ServicioViewSet(viewsets.ModelViewSet):
    queryset = Servicio.objects.all()
    serializer_class = ServicioSerializer

class TallerServicioViewSet(viewsets.ModelViewSet):
    queryset = TallerServicio.objects.all()
    serializer_class = TallerServicioSerializer

class FavoritoTallerViewSet(viewsets.ModelViewSet):
    queryset = FavoritoTaller.objects.all()
    serializer_class = FavoritoTallerSerializer

# ViewSets para modelos relacionados al vehículo
class MarcaViewSet(viewsets.ModelViewSet):
    queryset = Marca.objects.all()
    serializer_class = MarcaSerializer

class TipoVehiculoViewSet(viewsets.ModelViewSet):
    queryset = TipoVehiculo.objects.all()
    serializer_class = TipoVehiculoSerializer

class VehiculoViewSet(viewsets.ModelViewSet):
    queryset = Vehiculo.objects.all()
    serializer_class = VehiculoSerializer

# ViewSets para modelos relacionados a la agenda
class EstadoAgendaViewSet(viewsets.ModelViewSet):
    queryset = EstadoAgenda.objects.all()
    serializer_class = EstadoAgendaSerializer

class AgendaViewSet(viewsets.ModelViewSet):
    queryset = Agenda.objects.all()
    serializer_class = AgendaSerializer

# ViewSets para modelos relacionados a los pagos
class FormaPagoViewSet(viewsets.ModelViewSet):
    queryset = FormaPago.objects.all()
    serializer_class = FormaPagoSerializer

class ReportePagoViewSet(viewsets.ModelViewSet):
    queryset = ReportePago.objects.all()
    serializer_class = ReportePagoSerializer

# ViewSets para modelos relacionados a los tickets
class EstadoTicketViewSet(viewsets.ModelViewSet):
    queryset = EstadoTicket.objects.all()
    serializer_class = EstadoTicketSerializer

class TicketViewSet(viewsets.ModelViewSet):
    queryset = Ticket.objects.all()
    serializer_class = TicketSerializer

# ViewSets para el modelo de contacto
class ContactoViewSet(viewsets.ModelViewSet):
    queryset = Contacto.objects.all()
    serializer_class = ContactoSerializer

# Eliminar servicios del taller desde la api
@api_view(['DELETE'])
def eliminar_servicios_de_taller(request, idTaller):
    try:
        qs = TallerServicio.objects.filter(idTaller=idTaller)

        if not qs.exists():
            return Response(
                {"detail": "El taller no tiene servicios."},
                status=status.HTTP_200_OK
            )

        qs.delete()

        return Response(
            {"detail": "Servicios eliminados correctamente."},
            status=status.HTTP_200_OK
        )
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

def index(request):
	return render(request, 'core/index.html')

@role_required(["Administrador"])
def administrar_talleres(request):
    talleres = Taller.objects.all()
    return render(request, 'core/administrar_talleres.html', {'talleres': talleres})
    

@role_required(["Administrador"])
def crear_taller(request):
    form = TallerForm()
    coordinates = {}
    address = ""

    if request.method == 'POST':
        # ----- Crear taller -----
        if 'crear_taller' in request.POST:
            form = TallerForm(request.POST, request.FILES)
            if form.is_valid():
                taller = form.save(commit=False)
                latitud = request.POST.get('latitud')
                longitud = request.POST.get('longitud')
                direccion = request.POST.get('direccion')

                if latitud and longitud:
                    taller.latitud = float(latitud)
                    taller.longitud = float(longitud)
                if direccion:
                    taller.direccion = direccion
                taller.save()

                servicios_seleccionados = form.cleaned_data.get('servicios', [])
                for servicio in servicios_seleccionados:
                    TallerServicio.objects.create(idTaller=taller, idServicio=servicio)

                return redirect('administrar_talleres')
            else:
                print("hola soy un error en la vista de crear_taller :D")

        # ----- Obtener ubicación -----
        elif 'obtener_ubicacion' in request.POST:
            address = request.POST.get('address', '').strip()
            if not address:
                return JsonResponse({'success': False, 'error': 'No se ingresó dirección.'})

            try:
                encoded_address = quote_plus(address)
                url = f'https://nominatim.openstreetmap.org/search?q={encoded_address}&format=json&addressdetails=1&limit=1'
                headers = {'User-Agent': 'FixSpotApp (contacto: soporte@fixspot.cl)'}
                response = requests.get(url, headers=headers, timeout=10)
                response.raise_for_status()
                data = response.json()

                if not data:
                    return JsonResponse({'success': False, 'error': 'No se encontró la dirección especificada.'})

                location = data[0]
                lat = location.get('lat')
                lon = location.get('lon')

                if not lat or not lon:
                    return JsonResponse({'success': False, 'error': 'No se pudieron obtener coordenadas.'})

                coordinates = {'lat': float(lat), 'lng': float(lon)}
                return JsonResponse({'success': True, 'coordinates': coordinates, 'address': address})

            except requests.RequestException as e:
                print(f"[ERROR NOMINATIM REQUEST] {str(e)}")
                return JsonResponse({'success': False, 'error': 'Error al conectarse con el servicio de geolocalización.'})
            except ValueError as e:
                print(f"[ERROR JSON PARSE] {str(e)}")
                return JsonResponse({'success': False, 'error': 'Error al procesar la respuesta del servidor.'})

    return render(request, 'core/crear_taller.html', {'form': form})


@role_required(["Administrador"])
def modificar_taller(request, id_taller):
    taller = get_object_or_404(Taller, idTaller=id_taller)

    if request.method == 'POST':
        # ----- Modificar taller -----
        if 'modificar_taller' in request.POST:
            form = TallerForm(request.POST, request.FILES, instance=taller)
            if form.is_valid():
                taller_modificado = form.save(commit=False)
                if 'direccion' in request.POST:
                    taller_modificado.direccion = request.POST['direccion']
                taller_modificado.save()

                TallerServicio.objects.filter(idTaller=taller_modificado).delete()
                servicios_seleccionados = form.cleaned_data.get('servicios', [])
                for servicio in servicios_seleccionados:
                    TallerServicio.objects.create(idTaller=taller_modificado, idServicio=servicio)

                return redirect('administrar_talleres')
            else:
                print("Error al modificar el taller. Revisa los campos.")

        # ----- Obtener ubicación -----
        elif 'obtener_ubicacion' in request.POST:
            address = request.POST.get('address', '').strip()
            if not address:
                return JsonResponse({'success': False, 'error': 'No se ingresó dirección.'})

            try:
                encoded_address = quote_plus(address)
                url = f'https://nominatim.openstreetmap.org/search?q={encoded_address}&format=json&addressdetails=1&limit=1'
                headers = {'User-Agent': 'FixSpotApp (contacto: soporte@fixspot.cl)'}
                response = requests.get(url, headers=headers, timeout=10)
                response.raise_for_status()
                data = response.json()

                if not data:
                    return JsonResponse({'success': False, 'error': 'No se encontró la dirección especificada.'})

                location = data[0]
                lat = location.get('lat')
                lon = location.get('lon')

                if not lat or not lon:
                    return JsonResponse({'success': False, 'error': 'No se pudieron obtener coordenadas.'})

                taller.direccion = address
                taller.save()

                coordinates = {'lat': float(lat), 'lng': float(lon)}
                return JsonResponse({'success': True, 'coordinates': coordinates, 'address': address})

            except requests.RequestException as e:
                print(f"[ERROR NOMINATIM REQUEST] {str(e)}")
                return JsonResponse({'success': False, 'error': 'Error al conectarse con el servicio de geolocalización.'})
            except ValueError as e:
                print(f"[ERROR JSON PARSE] {str(e)}")
                return JsonResponse({'success': False, 'error': 'Error al procesar la respuesta del servidor.'})

    else:
        form = TallerForm(instance=taller)

    return render(request, 'core/modificar_taller.html', {'form': form, 'taller': taller})

@role_required(["Administrador"])
def eliminar_taller(request, id_taller):
    taller = get_object_or_404(Taller, idTaller=id_taller)
    taller.delete()
    return redirect('administrar_talleres')

@login_required
def agendar_hora(request):
	return render(request, 'core/agendar_hora.html')

@login_required
def agendar(request):
	return render(request, 'core/agendar.html')

@login_required
def contactanos(request):
    if request.method == 'POST':
        form = ContactoForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('contactanos')
        else:
           pass
    else:
        form = ContactoForm()
    return render(request, 'core/contactanos.html', {'form': form})

def login(request):
    form = LoginForm()

    if request.method == "POST":
        form = LoginForm(request.POST)

        if form.is_valid():
            username = form.cleaned_data["username"].strip()
            password = form.cleaned_data["password"]

            if username == "" or password == "":
                messages.error(request, "Debes completar todos los campos.")
                return render(request, "registration/login.html", {"form": form})

            try:
                user = UsuarioCustom.objects.get(username=username)
            except UsuarioCustom.DoesNotExist:
                messages.error(request, "El usuario no existe.")
                return render(request, "registration/login.html", {"form": form})

            if not check_password(password, user.password):
                messages.error(request, "Contraseña incorrecta.")
                return render(request, "registration/login.html", {"form": form})

            user = authenticate(username=username, password=password)

            if user is None:
                messages.error(request, "Error al autenticar usuario.")
                return render(request, "registration/login.html", {"form": form})

            auth_login(request, user)

            # Mensaje de éxito SweetAlert2
            messages.success(request, "Sesión iniciada correctamente.")

            return redirect("index")

    return render(request, "registration/login.html", {"form": form})

class CustomLogoutView(LogoutView):
    next_page = 'index'  

    def dispatch(self, request, *args, **kwargs):
        messages.success(request, "Sesión cerrada correctamente.")
        return super().dispatch(request, *args, **kwargs)

@login_required
def mapa(request):
    talleres = Taller.objects.all()
    talleres_json = serialize('json', talleres, fields=('nombreTaller', 'direccion', 'latitud', 'longitud'))
    return render(request, 'core/mapa.html', {'talleres_json': talleres_json})

@login_required
def mis_reservas(request):
    agendas = Agenda.objects.filter(cliente=request.user)
    return render(request, 'core/mis_reservas.html', {'agendas': agendas})

@csrf_exempt
def actualizar_estado_agenda(request, id_agenda):
    if request.method == 'POST':
        agenda = get_object_or_404(Agenda, idAgenda=id_agenda)
        
        # Obtener la instancia de EstadoAgenda correspondiente
        try:
            estado_pagado = EstadoAgenda.objects.get(nombreEstado='Pagado')
        except EstadoAgenda.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Estado "Pagado" no encontrado'}, status=400)

        agenda.estado = estado_pagado
        agenda.save()

        # Devolver una respuesta JSON con success
        return JsonResponse({'success': True, 'filename': f'Reserva_{agenda.idAgenda}.docx'})

    return JsonResponse({'success': False}, status=400)

def generar_documento_word(request, id_agenda):
    agenda = get_object_or_404(Agenda, idAgenda=id_agenda)

    # Generar el documento Word
    doc = docx.Document()
    doc.add_heading('Detalle de la Reserva', level=1)
    doc.add_paragraph(f'Fecha de Atención: {agenda.fechaAtencion}')
    doc.add_paragraph(f'Hora de Atención: {agenda.horaAtencion}')
    doc.add_paragraph(f'Tipo de Agenda: {agenda.idServicio}')
    doc.add_paragraph(f'Taller: {agenda.idTaller}')
    doc.add_paragraph(f'Patente Vehículo: {agenda.idVehiculo}')

    # Detalles del Reporte de Pago
    reporte_pago = ReportePago.objects.filter(reserva=agenda).first()
    if reporte_pago:
        doc.add_heading('Detalle del Reporte de Pago', level=2)
        doc.add_paragraph(f'Comentario: {reporte_pago.comentario}')
        doc.add_paragraph(f'Monto a pagar: ${reporte_pago.monto}')

    # Guardar el documento en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Reserva_{agenda.idAgenda}.docx'
    return response

def generar_boleta_pdf(request, id_agenda):
    agenda = get_object_or_404(Agenda, idAgenda=id_agenda)
    reporte_pago = ReportePago.objects.filter(reserva=agenda).first()

    # Crear PDF
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # === ENCABEZADO ===
    p.setFillColorRGB(0.35, 0.16, 0.51)  # Morado FixSpot
    p.rect(0, height - 100, width, 100, fill=True, stroke=False)

    p.setFillColor(colors.white)
    p.setFont("Helvetica-Bold", 20)
    p.drawString(50, height - 60, "FIXSPOT")
    p.setFont("Helvetica", 12)
    p.drawString(50, height - 80, "Boleta de Servicio")
    p.setFont("Helvetica", 10)
    p.drawString(450, height - 60, f"Fecha emisión: {datetime.now().strftime('%d/%m/%Y')}")

    # === DATOS PRINCIPALES ===
    p.setFillColor(colors.black)
    p.setFont("Helvetica-Bold", 13)
    p.drawString(50, height - 130, "DETALLE DE LA RESERVA")
    p.setFont("Helvetica", 11)
    p.drawString(50, height - 150, f"Taller: {agenda.idTaller.nombreTaller}")
    p.drawString(50, height - 165, f"Servicio: {agenda.idServicio.nombreServicio}")
    p.drawString(50, height - 180, f"Cliente: {agenda.cliente.pnombre} {agenda.cliente.ap_paterno}")
    p.drawString(50, height - 195, f"Vehículo: {agenda.idVehiculo.patente}")
    p.drawString(50, height - 210, f"Fecha atención: {agenda.fechaAtencion.strftime('%d/%m/%Y')}  |  Hora: {agenda.horaAtencion.strftime('%H:%M')}")

    # === LÍNEA DIVISORIA ===
    p.setStrokeColor(colors.grey)
    p.line(50, height - 225, width - 50, height - 225)

    # === DETALLE DEL PAGO ===
    y = height - 260
    if reporte_pago:
        p.setFont("Helvetica-Bold", 13)
        p.drawString(50, y, "DETALLE DEL PAGO")
        y -= 20
        p.setFont("Helvetica", 11)
        metodo = reporte_pago.idFormaPago.nombreFormaPago if hasattr(reporte_pago, 'idFormaPago') and reporte_pago.idFormaPago else "Efectivo"
        p.drawString(50, y, f"Método de pago: {metodo}")
        y -= 15
        p.drawString(50, y, f"Comentario: {reporte_pago.comentario or 'N/A'}")
        y -= 15
        p.drawString(50, y, f"Monto Total: ${reporte_pago.monto:,.0f}")

    # === SECCIÓN TOTAL ===
    y -= 40
    p.setFillColorRGB(0.95, 0.95, 0.95)
    p.roundRect(50, y - 40, width - 100, 50, 8, fill=True, stroke=False)
    p.setFillColor(colors.black)
    p.setFont("Helvetica-Bold", 12)
    p.drawString(60, y - 15, "Total a Pagar:")
    p.setFont("Helvetica-Bold", 14)
    p.drawRightString(width - 70, y - 15, f"${reporte_pago.monto:,.0f}")

    # === FOOTER ===
    p.setFont("Helvetica-Oblique", 9)
    p.setFillColor(colors.grey)
    p.drawString(50, 80, "Documento generado automáticamente por FixSpot")
    p.drawString(50, 65, "Este documento no requiere firma ni timbre.")
    p.drawString(50, 50, "© FixSpot - Todos los derechos reservados.")

    # Finalizar PDF
    p.showPage()
    p.save()
    buffer.seek(0)

    # Nombre del archivo con formato claro
    filename = f"Boleta_FixSpot_{agenda.idTaller.nombreTaller}_{agenda.idAgenda}.pdf".replace(" ", "_")

    return HttpResponse(buffer, content_type="application/pdf", headers={
        "Content-Disposition": f'attachment; filename="{filename}"'
    })


def detalle_reserva(request, id_agenda):
    agenda = get_object_or_404(Agenda, idAgenda=id_agenda)
    reporte_pago = ReportePago.objects.filter(reserva=agenda).first()
    formas_pago = FormaPago.objects.all()

    try:
        respuesta = requests.get('https://mindicador.cl/api/')
        monedas = respuesta.json()
        tasa_dolar = Decimal(monedas['dolar']['valor'])
    except requests.exceptions.RequestException:
        tasa_dolar = None

    monto_dolares = None
    if reporte_pago and tasa_dolar:
        monto_dolares = round(Decimal(reporte_pago.monto) / tasa_dolar, 2)
        monto_dolares = format(monto_dolares, '.2f')

    context = {
        'agenda': agenda,
        'reporte_pago': reporte_pago,
        'monto_dolares': monto_dolares,
        'formas_pago': formas_pago,
    }
    return render(request, 'core/detalle_reserva.html', context)

@login_required
def mis_vehiculos(request):
    vehiculos = Vehiculo.objects.filter(idUsuario=request.user)
    return render(request, 'core/mis_vehiculos.html', {'vehiculos': vehiculos})

@login_required
@role_required(["Encargado taller"])
def realizar_ticket(request):
	return render(request, 'core/realizar_ticket.html')

@login_required
def solicitar_taller(request):
	return render(request, 'core/solicitar_taller.html')

@login_required
def register_vehiculo(request):
	return render(request, 'core/register_vehiculo.html')

def register(request):
    form = RegisterForm()
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)

            user.idRol = rolUsuario.objects.get(nombreRol='Cliente')
            user.idComuna = form.cleaned_data['comuna']
            user.email = user.email
            
            user.save()
            messages.success(request, '¡Te has registrado con éxito!')

            return redirect("index")

    return render(request, 'registration/register.html', {'form': form})

def sobre_nosotros(request):
	return render(request, 'core/sobre_nosotros.html')

@login_required
def talleres(request):
    base = Taller.objects.annotate(
        promedio_calificacion=Avg('calificaciontaller__calificacion')
    )
    if request.user.is_authenticated:
        fav_sub = FavoritoTaller.objects.filter(
            usuario=request.user, taller=OuterRef('pk')
        )
        base = base.annotate(es_favorito=Exists(fav_sub))
    else:
        base = base.annotate(es_favorito=models.Value(False, output_field=models.BooleanField()))
    return render(request, 'core/talleres.html', {'talleres': base})

@login_required
@role_required(["Administrador"])
def tickets(request):
    tickets = Ticket.objects.all()
    data = {
        'tickets': tickets
    }
    return render(request, 'core/tickets.html', data)

def get_coordinates(request):
    try:
        error_message=""
        success_message=""
        if request.method == 'POST':
            form = AddressForm(request.POST)
            if form.is_valid():
                address = form.cleaned_data['address']
                api_key = '855d3348e0ee4298b17b00e020377da5'
                base_url = 'https://api.opencagedata.com/geocode/v1/json'
                params = {'q': address, 'key': api_key}
                response = requests.get(base_url, params=params)
                if response.status_code == 200:
                    data = response.json()
                    print(data) 
                    if data['results']:
                        location = data['results'][0]['geometry']
                        address_details = data['results'][0]['components']

                        # Verificar si la dirección está en la Región Metropolitana
                        if 'state' in address_details and address_details['state'] == 'Santiago Metropolitan Region':
                            coordinates = {
                                'lat': location['lat'],
                                'lng': location['lng']
                            }
                            success_message = "La dirección seleccionada es válida"
                            return render(request, 'core/test.html', {'form': form, 'coordinates': coordinates, 'success_message':success_message})
                        else:
                            error_message = "La dirección seleccionada no está en la Región Metropolitana."
                            return render(request, 'core/test.html', {'form': form, 'error_message': error_message})
    except Exception as e:
        print(e)
    form = AddressForm()
    return render(request, 'core/test.html', {'form': form})

def autocomplete_address(request):
    if 'term' in request.GET:
        search_term = request.GET.get('term')
        url = 'https://nominatim.openstreetmap.org/search'
        params = {
            'q': search_term,
            'format': 'json',
            'addressdetails': 1,
            'limit': 5,
        }
        headers = {
            'User-Agent': 'FixSpotApp (contacto: soporte@fixspot.cl)'
        }
        try:
            response = requests.get(url, params=params, headers=headers, timeout=10)
            response.raise_for_status()
            suggestions = response.json()

            results = []
            for s in suggestions:
                display = s.get('display_name', '')
                if display:
                    results.append({'label': display, 'value': display})
            return JsonResponse(results, safe=False)
        except requests.RequestException as e:
            print(f"[AUTOCOMPLETE ERROR] {str(e)}")
            return JsonResponse([], safe=False)
    return JsonResponse([], safe=False)

@login_required
@role_required(["Administrador"])
def administracion(request):
    return render(request, 'core/administracion.html')

@login_required
def agendar_hora(request, id_taller):
    taller = get_object_or_404(Taller, idTaller=id_taller)

    if request.method == 'POST':
        form = AgendaForm(request.POST or None, user=request.user, taller=taller)
        if form.is_valid():
            agenda = form.save(commit=False)
            agenda.idTaller = taller
            agenda.cliente = request.user
            agenda.save()
            return redirect('mis_reservas')
    else:
        form = AgendaForm(user=request.user, taller=taller)

    return render(request, 'core/agendar_hora.html', {'form': form, 'taller': taller})


def get_available_hours(request):
    fecha = request.GET.get('date')  # Fecha seleccionada
    id_taller = request.GET.get('id_taller')  # Id del taller
    taller = get_object_or_404(Taller, idTaller=id_taller)

    # Generar las horas entre 9:00 AM y 7:00 PM, excluyendo las 14:00 PM (2:00 PM)
    available_hours = [f"{h:02}:00" for h in range(9, 19) if h != 14]

    if fecha:
        # Excluir horas ya reservadas para la fecha seleccionada
        reserved_hours = Agenda.objects.filter(
            fechaAtencion=fecha,
            idTaller=taller
        ).values_list('horaAtencion', flat=True)

        # Convertir las horas reservadas a formato de cadena ("HH:MM")
        reserved_hours = [hour.strftime("%H:%M") for hour in reserved_hours]  # Convertir de `time` a cadena

        print(f"Reserved hours for {fecha}: {reserved_hours}")  # Verificar las horas reservadas

        # Excluir las horas reservadas de la lista de horas disponibles
        available_hours = [hour for hour in available_hours if hour not in reserved_hours]

    print(f"Available hours after filtering: {available_hours}")  # Verificar las horas disponibles
    return JsonResponse({'available_hours': available_hours})


@login_required
def annadir_vehiculo(request):
    form = VehiculoForm()  

    if request.method == 'POST':
        form = VehiculoForm(request.POST)  

        if form.is_valid():
            vehiculo = form.save(commit=False)
            vehiculo.idUsuario = request.user  
            vehiculo.save()

            return redirect('mis_vehiculos')

    return render(request, 'core/annadir_vehiculo.html', {'form': form})

@login_required
def modificar_vehiculo(request, vehiculo_id):
    vehiculo = Vehiculo.objects.get(idVehiculo=vehiculo_id)

    if request.method == 'POST':
        form = VehiculoForm(request.POST, instance=vehiculo)
        if form.is_valid():
            form.save() 
            return redirect('mis_vehiculos')  
    else:
        form = VehiculoForm(instance=vehiculo)

    return render(request, 'core/modificar_vehiculo.html', {'form': form})

@login_required
def eliminar_vehiculo(request, vehiculo_id):
    vehiculo = Vehiculo.objects.get(idVehiculo = vehiculo_id)
    vehiculo.delete()
    return redirect('mis_vehiculos')
   
@login_required
@role_required(["Administrador"])
def administrar_usuarios(request):
    usuarios = UsuarioCustom.objects.all()
    return render(request, 'core/administrar_usuarios.html', {'usuarios': usuarios})

@login_required
@role_required(["Administrador"])
def modificar_usuario(request, id):
    usuario = get_object_or_404(UsuarioCustom, id=id)
    if request.method == 'POST':
        form = UsuarioCustomForm(request.POST, instance=usuario)
        if form.is_valid():
            form.save()
            return redirect('administrar_usuarios')
    else:
        form = UsuarioCustomForm(instance=usuario)
    return render(request, 'core/modificar_usuario.html', {'form': form})

@login_required 
@role_required(["Administrador"])
def eliminar_usuario(request, id):
    usuario = UsuarioCustom.objects.get(id = id)
    if request.method == 'POST':
        usuario.delete()
        return redirect('administrar_usuarios')

@login_required 
@role_required(["Administrador"])
def crear_usuario(request):
    if request.method == 'POST':
        form = UsuarioCustomCreationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('administrar_usuarios')
    else:
        form = UsuarioCustomCreationForm()
    return render(request, 'core/crear_usuario.html', {'form': form})

@login_required
@role_required(["Encargado taller"])
def administrar_mi_taller(request):
    taller_del_usuario = Taller.objects.filter(idUsuario=request.user).first()

    return render(request, 'core/administrar_mi_taller.html', {'taller': taller_del_usuario})

@login_required
@role_required(["Encargado taller"])
def reservas_taller(request, idTaller):
    taller = get_object_or_404(Taller, pk=idTaller)
    reservas = (Agenda.objects
                .filter(idTaller=taller)
                .select_related('idServicio', 'cliente', 'idVehiculo__idMarca', 'estado')
               )
    hay_pagadas = reservas.filter(estado__idEstado=3).exists()
    data = {
        'taller': taller,
        'reservas': reservas,
        'hay_pagadas': hay_pagadas,
    }
    return render(request, 'core/reservas_taller.html', data)

@login_required
def perfil_usuario(request):
    user = request.user
    pref, _ = PreferenciasUsuario.objects.get_or_create(usuario=user)

    if request.method == 'POST':
        form = UsuarioCustomPerfilForm(request.POST, instance=user)
        # --- guardar primero usuario ---
        if form.is_valid():
            form.save()
            # --- guardar toggle acepta_promociones ---
            raw = request.POST.get('acepta_promociones', '') 
            acepta = str(raw).lower() in ('1', 'true', 'on', 'yes', 'si', 'sí')
            if pref.acepta_promociones != acepta:
                pref.acepta_promociones = acepta
                pref.save()

            messages.success(request, 'Perfil actualizado correctamente!')
            return redirect('index')
    else:
        form = UsuarioCustomPerfilForm(instance=user)

    ctx = {
        'form': form,
        'pref_acepta_promos': pref.acepta_promociones,
    }
    return render(request, 'core/perfil_usuario.html', ctx)

@role_required(["Encargado taller"])
def generar_reporte_pago(request, idReserva):
    reserva = get_object_or_404(Agenda, pk=idReserva)
    if request.method == 'POST':
        form = ReportePagoForm(request.POST)
        if form.is_valid():
            reporte = form.save(commit=False)
            reporte.reserva = reserva
            reporte.save()

            reserva.estado_id = 2  
            reserva.save()

            return redirect('reservas_taller', idTaller=reserva.idTaller.idTaller)
    else:
        form = ReportePagoForm()
    
    return render(request, 'core/generar_reporte_pago.html', {'form': form, 'reserva': reserva})

@login_required 
@role_required(["Encargado taller"])
def mis_tickets(request):
    user = request.user
    tickets_usuario = Ticket.objects.filter(solicitante=user)
    data = {
        'tickets_usuario':tickets_usuario
    }
    return render(request, 'core/mis_tickets.html', data)

@login_required
@role_required(["Encargado taller"])
def crear_ticket(request):
    if request.method == 'POST':
        form = TicketForm(request.POST)
        if form.is_valid():
            ticket = form.save(commit=False)
            ticket.solicitante = request.user
            estado_pendiente = EstadoTicket.objects.get(NombreEstado='Pendiente')
            ticket.EstadoTicket = estado_pendiente
            ticket.save()
            return redirect('mis_tickets') 
    else:
        form = TicketForm()
    return render(request, 'core/crear_ticket.html', {'form': form})

def aceptar_ticket(request, idTicket):
    ticket = get_object_or_404(Ticket, idTicket=idTicket)
    ticket.EstadoTicket = EstadoTicket.objects.get(NombreEstado='Aceptado')
    ticket.save()
    return redirect('tickets')

def rechazar_ticket(request, idTicket):
    ticket = get_object_or_404(Ticket, idTicket=idTicket)
    ticket.EstadoTicket = EstadoTicket.objects.get(NombreEstado='Rechazado')
    ticket.save()
    return redirect('tickets')

def mensajes(request):
    mensajes = Contacto.objects.all()
    return render(request, 'core/mensajes.html', {'mensajes': mensajes})

def eliminar_mensaje(request, id_mensaje):
    mensaje = get_object_or_404(Contacto, idContacto=id_mensaje)
    mensaje.delete()
    return redirect('mensajes')

def guardar_calificacion(request, agenda_id):
    if request.method == 'POST':
        calificacion = request.POST.get('calificacion')
        agenda = get_object_or_404(Agenda, idAgenda=agenda_id)

        # Crear la calificación
        CalificacionTaller.objects.create(
            idTaller=agenda.idTaller,
            idUsuario=agenda.cliente,
            calificacion=calificacion
        )

        return JsonResponse({'success': True})
    return JsonResponse({'success': False}, status=400)


#VISTAS TALLERES FAVORITOS
@login_required
def favoritos_list(request):
    favoritos = (FavoritoTaller.objects
                 .filter(usuario=request.user)
                 .select_related('taller'))
    return render(request, 'core/favoritos_list.html', {'favoritos': favoritos})

@login_required
def toggle_favorito(request):
    if request.method != 'POST':
        return HttpResponseBadRequest('Método no permitido')
    taller_id = request.POST.get('taller_id')
    if not taller_id:
        return HttpResponseBadRequest('Falta taller_id')

    taller = get_object_or_404(Taller, pk=taller_id)
    fav_qs = FavoritoTaller.objects.filter(usuario=request.user, taller=taller)

    # si existe → quitar
    if fav_qs.exists():
        fav_qs.delete()
        return JsonResponse({'ok': True, 'favorito': False})

    # validar límite 5
    if FavoritoTaller.objects.filter(usuario=request.user).count() >= 5:
        return JsonResponse({'ok': False, 'error': 'max', 'msg': 'Máximo 5 talleres favoritos'}, status=400)

    FavoritoTaller.objects.create(usuario=request.user, taller=taller)
    return JsonResponse({'ok': True, 'favorito': True})


def _fmt_clp(v):
    if v is None:
        return ""
    try:
        return f"${int(round(float(v))):,}".replace(",", ".")
    except Exception:
        return f"${v}"

@login_required
def export_taller_pagadas_excel(request, idTaller):
    """
    Exporta un Excel con TODAS las reservas PAGADAS (estado id=3) del taller,
    usando el monto real desde ReportePago.monto. Sin IVA, solo 'Precio total'.
    """
    taller = get_object_or_404(Taller, pk=idTaller)

    # Trae reservas pagadas y su ReportePago (OneToOne reverse: 'reportepago')
    reservas = (Agenda.objects
                .filter(idTaller=taller, estado__idEstado=3)
                .select_related('idServicio', 'cliente', 'idVehiculo__idMarca', 'estado', 'reportepago')
                .order_by('fechaAtencion', 'horaAtencion', 'idAgenda'))

    if not reservas.exists():
        return HttpResponseForbidden("No hay reservas pagadas para exportar.")

    # ===== Excel =====
    wb = Workbook()
    ws = wb.active
    ws.title = "Pagadas"

    # Estilos
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    title_font = Font(size=16, bold=True, color="2E2E2E")
    header_font = Font(size=11, bold=True, color="1F1F1F")
    fill_header = PatternFill("solid", fgColor="F2F2F2")

    # Anchos de columnas: ID, Fecha, Hora, Servicio, Cliente, Patente, Vehículo, Precio total
    widths = [7, 14, 8, 28, 26, 16, 28, 16]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Título
    ws.merge_cells("A1:H1")
    ws["A1"] = f"Reservas Pagadas — {taller.nombreTaller.upper()}"
    ws["A1"].font = title_font
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # Encabezados
    headers = ["ID", "Fecha", "Hora", "Servicio", "Cliente", "Patente", "Vehículo", "Precio total"]
    row = 3
    for col, h in enumerate(headers, start=1):
        c = ws.cell(row=row, column=col, value=h)
        c.font = header_font
        c.fill = fill_header
        c.alignment = Alignment(horizontal="center")
        c.border = border

    # Filas
    total_general = 0.0
    row += 1
    for r in reservas:
        # monto desde ReportePago (si por cualquier razón no existe, 0)
        monto = float(getattr(getattr(r, "reportepago", None), "monto", 0) or 0)
        total_general += monto

        veh_marca = getattr(r.idVehiculo.idMarca, 'nombreMarca', '')
        veh_modelo = r.idVehiculo.modelo or ''
        veh_sub = f" {r.idVehiculo.subModelo}" if r.idVehiculo.subModelo else ''
        veh_text = f"{veh_marca} {veh_modelo}{veh_sub}".strip()

        ws.cell(row=row, column=1, value=f"#{r.idAgenda}").alignment = Alignment(horizontal="center")
        ws.cell(row=row, column=2, value=str(r.fechaAtencion)).alignment = Alignment(horizontal="center")
        ws.cell(row=row, column=3, value=str(r.horaAtencion)).alignment = Alignment(horizontal="center")
        ws.cell(row=row, column=4, value=str(r.idServicio)).alignment = Alignment(horizontal="left")
        ws.cell(row=row, column=5, value=f"{r.cliente.pnombre} {r.cliente.ap_paterno}").alignment = Alignment(horizontal="left")
        ws.cell(row=row, column=6, value=r.idVehiculo.patente).alignment = Alignment(horizontal="center")
        ws.cell(row=row, column=7, value=veh_text).alignment = Alignment(horizontal="left")
        ws.cell(row=row, column=8, value=_fmt_clp(monto)).alignment = Alignment(horizontal="right")

        for c in range(1, 9):
            ws.cell(row=row, column=c).border = border

        row += 1

    # Total general
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    lbl = ws.cell(row=row, column=1, value="Total general")
    lbl.font = Font(bold=True)
    lbl.alignment = Alignment(horizontal="right")
    tot_cell = ws.cell(row=row, column=8, value=_fmt_clp(total_general))
    tot_cell.alignment = Alignment(horizontal="right")
    tot_cell.font = Font(bold=True, color="1B5E20")
    for c in range(1, 9):
        ws.cell(row=row, column=c).border = border

    # Respuesta
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    fname = f"reservas_pagadas_{taller.nombreTaller.replace(' ', '_')}.xlsx"
    resp = HttpResponse(
        bio.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    resp["Content-Disposition"] = f'attachment; filename=\"{fname}\"'
    return resp

